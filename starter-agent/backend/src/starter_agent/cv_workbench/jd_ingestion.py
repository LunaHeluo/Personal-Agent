"""Safe JD file parsing for workbench candidates.

Files are parsed in memory and only the normalized Markdown is handed to the
existing candidate service.  This preserves its authorization, artifact
retention, confirmation, and immutable-snapshot boundaries.
"""

from __future__ import annotations

import io
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol

import httpx
from docx import Document
from pypdf import PdfReader

from starter_agent.cv_workbench.resume_profile import infer_resume_name


class JobDocumentError(RuntimeError):
    """A user-actionable JD upload or parsing error."""


@dataclass(frozen=True)
class ParsedJobDocument:
    filename: str
    markdown: str
    extraction_method: str


class OcrParser(Protocol):
    def parse(self, *, filename: str, content: bytes) -> str: ...


class MinerUOcrParser:
    """Small MinerU adapter for image and PDF JD sources."""

    def __init__(self, *, token: str, base_url: str, timeout_seconds: int = 120) -> None:
        self.token = token
        self.base_url = base_url.rstrip("/")
        self.timeout_seconds = timeout_seconds

    def parse(self, *, filename: str, content: bytes) -> str:
        if not self.token:
            raise JobDocumentError("jd_ocr_not_configured")
        headers = {"Authorization": f"Bearer {self.token}"}
        with httpx.Client(timeout=self.timeout_seconds) as client:
            requested = client.post(
                f"{self.base_url}/api/v4/file-urls/batch",
                headers=headers,
                json={"files": [{"name": filename}], "model_version": "vlm"},
            )
            self._require_ok(requested, "jd_ocr_upload_url_failed")
            data = requested.json().get("data") or {}
            batch_id, urls = data.get("batch_id"), data.get("file_urls") or []
            if not isinstance(batch_id, str) or not urls:
                raise JobDocumentError("jd_ocr_upload_url_invalid")
            uploaded = client.put(str(urls[0]), content=content)
            if uploaded.status_code not in {200, 201}:
                raise JobDocumentError("jd_ocr_upload_failed")
            deadline = time.monotonic() + self.timeout_seconds
            result: dict[str, object] | None = None
            while time.monotonic() < deadline:
                poll = client.get(
                    f"{self.base_url}/api/v4/extract-results/batch/{batch_id}",
                    headers=headers,
                )
                self._require_ok(poll, "jd_ocr_status_failed")
                results = (poll.json().get("data") or {}).get("extract_result") or []
                if results and all(item.get("state") in {"done", "failed"} for item in results):
                    result = next((item for item in results if item.get("state") == "done"), None)
                    break
                time.sleep(2)
            if result is None:
                raise JobDocumentError("jd_ocr_parse_failed")
            archive_url = result.get("full_zip_url")
            if not isinstance(archive_url, str):
                raise JobDocumentError("jd_ocr_result_missing")
            archive = client.get(archive_url)
            if archive.status_code != 200:
                raise JobDocumentError("jd_ocr_result_download_failed")
        try:
            with zipfile.ZipFile(io.BytesIO(archive.content)) as values:
                item = next(
                    (entry for entry in values.infolist() if Path(entry.filename).name == "full.md"),
                    None,
                )
                if item is None:
                    raise JobDocumentError("jd_ocr_markdown_missing")
                return values.read(item).decode("utf-8").strip()
        except zipfile.BadZipFile as error:
            raise JobDocumentError("jd_ocr_result_invalid") from error

    @staticmethod
    def _require_ok(response: httpx.Response, code: str) -> None:
        if response.status_code != 200:
            raise JobDocumentError(code)
        try:
            if response.json().get("code") != 0:
                raise JobDocumentError(code)
        except ValueError as error:
            raise JobDocumentError(code) from error


_DIRECT_TEXT_EXTENSIONS = frozenset({".txt", ".md", ".markdown"})
_OCR_EXTENSIONS = frozenset({".pdf", ".png", ".jpg", ".jpeg", ".webp"})


def _extract_local_pdf_text(content: bytes) -> str:
    """Extract selectable text from a PDF without sending it to a cloud OCR service."""
    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise JobDocumentError("resume_pdf_encrypted")
        text = "\n".join(page.extract_text(extraction_mode="layout") or "" for page in reader.pages).strip()
    except JobDocumentError:
        raise
    except Exception as error:
        raise JobDocumentError("resume_pdf_parse_failed") from error
    if not text:
        raise JobDocumentError("resume_pdf_text_unavailable")
    return text


def parse_job_document(
    *, filename: str, content: bytes, ocr: OcrParser | None = None, max_bytes: int = 8_000_000
) -> ParsedJobDocument:
    """Extract one JD document without persisting the unconfirmed source file."""
    extension = Path(filename).suffix.lower()
    if not filename or extension not in _DIRECT_TEXT_EXTENSIONS | _OCR_EXTENSIONS | {".docx"}:
        raise JobDocumentError("jd_file_type_unsupported")
    if not content:
        raise JobDocumentError("jd_file_empty")
    if len(content) > max_bytes:
        raise JobDocumentError("jd_file_too_large")
    if extension in _DIRECT_TEXT_EXTENSIONS:
        try:
            text = content.decode("utf-8-sig").strip()
        except UnicodeDecodeError as error:
            raise JobDocumentError("jd_text_encoding_invalid") from error
        method = "text"
    elif extension == ".docx":
        try:
            text = "\n".join(paragraph.text for paragraph in Document(io.BytesIO(content)).paragraphs).strip()
        except Exception as error:
            raise JobDocumentError("jd_docx_parse_failed") from error
        method = "docx"
    else:
        if ocr is None:
            raise JobDocumentError("jd_ocr_not_configured")
        text = ocr.parse(filename=filename, content=content).strip()
        method = "mineru_ocr"
    if not text:
        raise JobDocumentError("jd_content_empty_after_parse")
    return ParsedJobDocument(filename=filename, markdown=text, extraction_method=method)


def parse_resume_document(
    *, filename: str, content: bytes, ocr: OcrParser | None = None, max_bytes: int = 8_000_000
) -> ParsedJobDocument:
    """Extract a DOCX/PDF resume for the Markdown-based versioning pipeline."""
    extension = Path(filename).suffix.lower()
    if extension not in {".docx", ".pdf"}:
        raise JobDocumentError("resume_file_type_unsupported")
    if extension == ".docx":
        return parse_job_document(filename=filename, content=content, ocr=ocr, max_bytes=max_bytes)
    if not content:
        raise JobDocumentError("jd_file_empty")
    if len(content) > max_bytes:
        raise JobDocumentError("jd_file_too_large")

    # MinerU remains the preferred parser for complex visual layouts, but a
    # regular PDF should still import when cloud OCR is unavailable or fails.
    if ocr is not None:
        try:
            text = ocr.parse(filename=filename, content=content).strip()
            if text:
                return ParsedJobDocument(filename=filename, markdown=text, extraction_method="mineru_ocr")
        except JobDocumentError:
            pass
    return ParsedJobDocument(
        filename=filename,
        markdown=_extract_local_pdf_text(content),
        extraction_method="local_pdf",
    )
