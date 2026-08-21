"""Version-bound PDF/DOCX exports stored as restricted immutable artifacts."""

from __future__ import annotations

import base64
import io
import json
import os
import re
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from hashlib import sha256
from pathlib import Path
from typing import Literal
from urllib.parse import urlparse
from uuid import NAMESPACE_URL, uuid5

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import KeepTogether, ListFlowable, ListItem, Paragraph, SimpleDocTemplate

from starter_agent.cv_workbench.contracts import (
    BusinessOperation,
    ExportRecord,
    ExportStatus,
    OperationStatus,
    ResumeVersion,
    ResumeVersionStatus,
)
from starter_agent.cv_workbench.operations import (
    BusinessOperationService,
    CommitReceipt,
    OperationCheckpoint,
    OperationCommand,
    RunBinding,
    RunOutcome,
    SafetyDecision,
    ValidationDecision,
)
from starter_agent.cv_workbench.store import ObjectNotFoundError, SQLiteWorkbenchStore
from starter_agent.cv_workbench.version_adapters import SessionKnowledgeVersionContentRepository
from starter_agent.infrastructure.session_store import SQLiteSessionStore


class ExportServiceError(RuntimeError):
    code = "export_service_error"


class ExportVersionNotConfirmedError(ExportServiceError):
    code = "export_requires_confirmed_version"


class ExportArtifactUnavailableError(ExportServiceError):
    code = "export_artifact_unavailable"


@dataclass(frozen=True)
class ExportCommand:
    operation_id: str
    idempotency_key: str
    export_id: str
    workspace_id: str
    resume_version_id: str
    format: Literal["pdf", "docx"]
    template_id: str = "ats-clean"
    template_version: str = "1.0.0"
    settings: dict[str, object] | None = None


@dataclass(frozen=True)
class ExportResult:
    operation: BusinessOperation
    record: ExportRecord | None


@dataclass(frozen=True)
class ExportArtifact:
    artifact_id: str
    content: bytes
    content_sha256: str
    filename: str
    media_type: str
    metadata: dict[str, object]
    expired: bool = False


class RestrictedExportArtifactRepository:
    def __init__(self, artifacts: SQLiteSessionStore, *, retention_days: int = 7, clock=lambda: datetime.now(UTC)) -> None:
        self.artifacts = artifacts
        self.retention_days = retention_days
        self.clock = clock

    def write(self, *, export_id: str, operation_id: str, workspace_id: str, principal: str, filename: str, media_type: str, content: bytes, metadata: dict[str, object]) -> ExportArtifact:
        digest = sha256(content).hexdigest()
        source_ref = f"artifact:resume-export:{export_id}:{digest[:16]}"
        namespace = uuid5(NAMESPACE_URL, f"resume-export:{workspace_id}:{export_id}")
        summary = dict(metadata) | {
            "source_type": "resume_export",
            "filename": filename,
            "media_type": media_type,
            "byte_length": len(content),
            "complete": True,
        }
        self.artifacts.save_tool_artifact(
            source_ref=source_ref,
            session_id=namespace,
            turn_id=uuid5(namespace, operation_id),
            tool_name="resume_export",
            content=base64.b64encode(content).decode("ascii"),
            call_id=operation_id,
            content_sha256=digest,
            truncation_summary=summary,
            parent_run_id=f"local-export:{operation_id}",
            access_level="restricted",
            principal=principal,
            expires_at=self.clock() + timedelta(days=self.retention_days),
        )
        return ExportArtifact(source_ref, content, digest, filename, media_type, summary)

    def read(self, artifact_id: str, *, principal: str) -> ExportArtifact:
        value = self.artifacts.get_tool_artifact_for_principal(artifact_id, principal=principal)
        if value is None:
            raise ExportArtifactUnavailableError("export_artifact_not_found")
        summary = dict(value.get("truncation_summary") or {})
        if value.get("expired"):
            return ExportArtifact(
                artifact_id, b"", str(value.get("content_sha256") or ""),
                str(summary.get("filename") or "export.bin"),
                str(summary.get("media_type") or "application/octet-stream"), summary, True,
            )
        try:
            content = base64.b64decode(str(value.get("content") or ""), validate=True)
        except ValueError as exc:
            raise ExportArtifactUnavailableError("export_artifact_corrupt") from exc
        digest = sha256(content).hexdigest()
        if digest != value.get("content_sha256"):
            raise ExportArtifactUnavailableError("export_artifact_hash_mismatch")
        return ExportArtifact(
            artifact_id, content, digest,
            str(summary.get("filename") or "export.bin"),
            str(summary.get("media_type") or "application/octet-stream"), summary,
        )


_LINK = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")


def _safe_link_markup(text: str) -> str:
    from xml.sax.saxutils import escape
    pieces: list[str] = []
    start = 0
    for match in _LINK.finditer(text):
        pieces.append(escape(text[start:match.start()]))
        url = match.group(2)
        if urlparse(url).scheme in {"http", "https"}:
            pieces.append(f'<link href="{escape(url)}" color="#2563EB">{escape(match.group(1))}</link>')
        else:
            pieces.append(escape(match.group(0)))
        start = match.end()
    pieces.append(escape(text[start:]))
    return "".join(pieces)


def _markdown_blocks(markdown: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    bullets: list[str] = []
    paragraph: list[str] = []

    def flush() -> None:
        nonlocal bullets, paragraph
        if paragraph:
            blocks.append(("paragraph", " ".join(paragraph)))
            paragraph = []
        if bullets:
            blocks.extend(("bullet", item) for item in bullets)
            bullets = []

    for raw in markdown.replace("\r\n", "\n").split("\n"):
        line = raw.strip()
        if not line:
            flush(); continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            flush(); blocks.append((f"heading{len(heading.group(1))}", heading.group(2))); continue
        bullet = re.match(r"^[-*+]\s+(.+)$", line)
        if bullet:
            if paragraph: flush()
            bullets.append(bullet.group(1)); continue
        if bullets: flush()
        paragraph.append(line)
    flush()
    return blocks


class AtsCleanRenderer:
    template_id = "ats-clean"
    template_version = "1.0.0"
    page_margin_mm = 16
    body_font_size = 9.5
    accent_color = "#0F4C5C"

    @staticmethod
    def _cjk_font() -> tuple[str, str]:
        candidates = [
            os.getenv("RESUME_AGENT_CJK_FONT", ""),
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simsun.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/System/Library/Fonts/PingFang.ttc",
        ]
        for value in candidates:
            if value and Path(value).is_file():
                return "ResumeAgentCJK", value
        raise ExportServiceError("cjk_font_unavailable")

    def render_pdf(self, markdown: str, settings: dict[str, object]) -> bytes:
        font_name, font_path = self._cjk_font()
        if font_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(font_name, font_path, subfontIndex=0))
        output = io.BytesIO()
        document = SimpleDocTemplate(
            output, pagesize=A4, rightMargin=self.page_margin_mm * mm, leftMargin=self.page_margin_mm * mm,
            topMargin=14 * mm, bottomMargin=14 * mm,
            title=str(settings.get("title") or "Resume"), author="Resume-Agent",
        )
        base = getSampleStyleSheet()
        body = ParagraphStyle("ResumeBody", parent=base["BodyText"], fontName=font_name, fontSize=self.body_font_size, leading=self.body_font_size + 3.5, textColor=colors.HexColor("#111827"), spaceAfter=4)
        styles = {
            "heading1": ParagraphStyle("ResumeTitle", parent=body, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=8),
            "heading2": ParagraphStyle("ResumeH2", parent=body, fontSize=12, leading=15, textColor=colors.HexColor(self.accent_color), spaceBefore=7, spaceAfter=4, borderWidth=0, keepWithNext=True),
            "heading3": ParagraphStyle("ResumeH3", parent=body, fontSize=10.5, leading=14, textColor=colors.HexColor("#1F2937"), spaceBefore=5, spaceAfter=3, keepWithNext=True),
            "paragraph": body,
        }
        story: list[object] = []
        pending_bullets: list[str] = []

        def flush_bullets() -> None:
            if not pending_bullets: return
            items = [ListItem(Paragraph(_safe_link_markup(item), body), leftIndent=4) for item in pending_bullets]
            story.append(ListFlowable(items, bulletType="bullet", start="circle", leftIndent=12, bulletFontName=font_name, bulletFontSize=7, spaceAfter=3))
            pending_bullets.clear()

        for kind, text in _markdown_blocks(markdown):
            if kind == "bullet": pending_bullets.append(text); continue
            flush_bullets()
            node = Paragraph(_safe_link_markup(text), styles.get(kind, body))
            story.append(KeepTogether([node]) if kind.startswith("heading") else node)
        flush_bullets()
        if not story: story.append(Paragraph("Resume", styles["heading1"]))
        document.build(story)
        return output.getvalue()

    @staticmethod
    def _set_run_font(run, name: str, size: float, *, bold: bool = False, color: str = "111827") -> None:
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
        run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)

    def render_docx(self, markdown: str, settings: dict[str, object]) -> bytes:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.55); section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(self.page_margin_mm / 25.4); section.right_margin = Inches(self.page_margin_mm / 25.4)
        section.header_distance = Inches(0.3); section.footer_distance = Inches(0.3)
        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(self.body_font_size)
        normal.paragraph_format.space_after = Pt(4); normal.paragraph_format.line_spacing = 1.08
        style_tokens = {
            "heading1": ("Title", 18, False, "111827", 0, 7),
            "heading2": ("Heading 1", 12, True, self.accent_color.removeprefix("#"), 7, 3),
            "heading3": ("Heading 2", 10.5, True, "1F2937", 5, 2),
        }
        for _kind, (name, size, bold, color, before, after) in style_tokens.items():
            style = doc.styles[name]
            style.font.name = "Microsoft YaHei"; style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(size); style.font.bold = bold; style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
        for kind, text in _markdown_blocks(markdown):
            style = style_tokens.get(kind, ("List Bullet" if kind == "bullet" else "Normal",))[0]
            paragraph = doc.add_paragraph(style=style)
            paragraph.paragraph_format.keep_together = True
            if kind == "heading1": paragraph.alignment = 1
            cursor = 0
            for match in _LINK.finditer(text):
                if match.start() > cursor:
                    self._set_run_font(paragraph.add_run(text[cursor:match.start()]), "Microsoft YaHei", self.body_font_size)
                self._add_hyperlink(paragraph, match.group(1), match.group(2))
                cursor = match.end()
            if cursor < len(text): self._set_run_font(paragraph.add_run(text[cursor:]), "Microsoft YaHei", self.body_font_size)
        properties = doc.core_properties
        properties.title = str(settings.get("title") or "Resume")
        properties.author = "Resume-Agent"
        output = io.BytesIO(); doc.save(output); return output.getvalue()

    @staticmethod
    def _add_hyperlink(paragraph, text: str, url: str) -> None:
        from docx.oxml import OxmlElement
        rel_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), rel_id)
        run = OxmlElement("w:r"); props = OxmlElement("w:rPr")
        color = OxmlElement("w:color"); color.set(qn("w:val"), "2563EB")
        underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
        fonts = OxmlElement("w:rFonts"); fonts.set(qn("w:eastAsia"), "Microsoft YaHei"); fonts.set(qn("w:ascii"), "Microsoft YaHei")
        props.extend((fonts, color, underline)); run.append(props)
        node = OxmlElement("w:t"); node.text = text; run.append(node); link.append(run); paragraph._p.append(link)

    def render(self, format: Literal["pdf", "docx"], markdown: str, settings: dict[str, object]) -> tuple[bytes, str]:
        if format == "pdf": return self.render_pdf(markdown, settings), "application/pdf"
        if format == "docx": return self.render_docx(markdown, settings), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        raise ExportServiceError("unsupported_export_format")


class AtsCompactRenderer(AtsCleanRenderer):
    template_id = "ats-compact"
    template_version = "1.0.0"
    page_margin_mm = 12
    body_font_size = 8.8
    accent_color = "#334155"


class ExportService:
    def __init__(self, *, store: SQLiteWorkbenchStore, content: SessionKnowledgeVersionContentRepository, artifacts: RestrictedExportArtifactRepository, renderer: AtsCleanRenderer | None = None, clock=lambda: datetime.now(UTC)) -> None:
        self.store = store; self.content = content; self.artifacts = artifacts
        selected = renderer or AtsCleanRenderer()
        self.renderer = selected  # Backward-compatible default renderer handle.
        self.renderers = {
            (selected.template_id, selected.template_version): selected,
            (AtsCompactRenderer.template_id, AtsCompactRenderer.template_version): AtsCompactRenderer(),
        }
        self.clock = clock

    def templates(self) -> tuple[dict[str, object], ...]:
        labels = {"ats-clean": "ATS 清爽", "ats-compact": "ATS 紧凑"}
        return tuple({"template_id": key[0], "template_version": key[1], "label": labels.get(key[0], key[0]), "formats": ("pdf", "docx")} for key in sorted(self.renderers))

    @staticmethod
    def _settings(command: ExportCommand) -> tuple[dict[str, object], str, str]:
        settings = dict(command.settings or {})
        canonical = json.dumps(settings, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
        settings_hash = sha256(canonical.encode()).hexdigest()
        input_hash = sha256("\0".join((command.resume_version_id, command.format, command.template_id, command.template_version, settings_hash)).encode()).hexdigest()
        return settings, settings_hash, input_hash

    def _find_reusable(self, command: ExportCommand, settings_hash: str, *, principal: str) -> ExportRecord | None:
        cursor = None
        while True:
            page = self.store.list(ExportRecord, principal=principal, cursor=cursor)
            for item in page.items:
                if (item.resume_version_id, item.format, item.template_id, item.template_version, item.settings_sha256) == (command.resume_version_id, command.format, command.template_id, command.template_version, settings_hash) and item.status == ExportStatus.AVAILABLE:
                    try:
                        artifact = self.artifacts.read(str(item.artifact_id), principal=principal)
                        if not artifact.expired: return item
                    except ExportArtifactUnavailableError:
                        pass
            if page.next_cursor is None: return None
            cursor = page.next_cursor

    def export(self, command: ExportCommand, *, principal: str) -> ExportResult:
        self.store.assert_entity_in_workspace(command.resume_version_id, command.workspace_id, principal=principal)
        version = self.store.get(ResumeVersion, command.resume_version_id, principal=principal)
        if version.status != ResumeVersionStatus.CONFIRMED:
            raise ExportVersionNotConfirmedError("export_requires_confirmed_version")
        renderer = self.renderers.get((command.template_id, command.template_version))
        if renderer is None:
            raise ExportServiceError("unsupported_export_template")
        settings, settings_hash, input_hash = self._settings(command)
        reusable = self._find_reusable(command, settings_hash, principal=principal)
        if reusable is not None:
            try:
                operation = self.store.get(BusinessOperation, command.operation_id, principal=principal)
            except ObjectNotFoundError:
                operation = self._commit_reuse_operation(command, input_hash, reusable, principal)
            return ExportResult(operation, reusable)
        committer = _ExportCommitter(self, command, settings_hash, principal)
        operations = BusinessOperationService(store=self.store, validator=_ExportValidator(self.artifacts, principal), safety_gate=_ExportSafetyGate(self.store, command, principal), committer=committer, clock=self.clock)
        operation, _ = operations.create(OperationCommand(command.operation_id, command.workspace_id, "export_resume", command.idempotency_key, input_hash), principal=principal)
        if operation.status == OperationStatus.COMMITTED:
            return ExportResult(operation, self.store.get(ExportRecord, str(operation.result_object_id), principal=principal))
        if operation.status == OperationStatus.COMMIT_FAILED:
            operation = operations.retry_commit(operation.operation_id, principal=principal)
            record = self.store.get(ExportRecord, str(operation.result_object_id), principal=principal) if operation.status == OperationStatus.COMMITTED else None
            return ExportResult(operation, record)
        run_id = f"local-export:{operation.operation_id}"
        if operation.status == OperationStatus.CREATED:
            operation = operations.bind_run(operation.operation_id, RunBinding(run_id), principal=principal)
        try:
            markdown = self.content.read(version.content, principal=principal, workspace_id=command.workspace_id)
            rendered, media_type = renderer.render(command.format, markdown, settings)
            artifact = self.artifacts.write(export_id=command.export_id, operation_id=command.operation_id, workspace_id=command.workspace_id, principal=principal, filename=f"{command.resume_version_id}.{command.format}", media_type=media_type, content=rendered, metadata={"export_id": command.export_id, "resume_version_id": command.resume_version_id, "format": command.format, "template_id": command.template_id, "template_version": command.template_version, "settings_sha256": settings_hash, "input_sha256": input_hash})
            operation = operations.process_run_outcome(operation.operation_id, RunOutcome(run_id, "succeeded", artifact.artifact_id, artifact.content_sha256), principal=principal)
        except Exception as exc:
            operations.process_run_outcome(operation.operation_id, RunOutcome(run_id, "failed", error_code=getattr(exc, "code", "export_render_failed")), principal=principal)
            raise
        record = self.store.get(ExportRecord, str(operation.result_object_id), principal=principal) if operation.status == OperationStatus.COMMITTED else None
        return ExportResult(operation, record)

    def _commit_reuse_operation(self, command: ExportCommand, input_hash: str, record: ExportRecord, principal: str) -> BusinessOperation:
        operations = BusinessOperationService(store=self.store, validator=_ReuseValidator(record), safety_gate=_ReuseSafetyGate(), committer=_ReuseCommitter(record), clock=self.clock)
        operation, _ = operations.create(OperationCommand(command.operation_id, command.workspace_id, "export_resume", command.idempotency_key, input_hash), principal=principal)
        if operation.status == OperationStatus.CREATED: operations.bind_run(operation.operation_id, RunBinding(f"local-export:{operation.operation_id}"), principal=principal)
        operation = operations.get(operation.operation_id, principal=principal)
        if operation.status == OperationStatus.RUNNING:
            operation = operations.process_run_outcome(operation.operation_id, RunOutcome(str(operation.parent_run_id), "succeeded", str(record.artifact_id), str(record.content_sha256)), principal=principal)
        return operation

    def download(self, export_id: str, *, principal: str) -> ExportArtifact:
        record = self.store.get(ExportRecord, export_id, principal=principal)
        if record.status != ExportStatus.AVAILABLE or not record.artifact_id: raise ExportArtifactUnavailableError("export_not_available")
        artifact = self.artifacts.read(record.artifact_id, principal=principal)
        if artifact.expired: raise ExportArtifactUnavailableError("export_download_expired")
        if artifact.content_sha256 != record.content_sha256: raise ExportArtifactUnavailableError("export_record_hash_mismatch")
        return artifact


class _ExportValidator:
    def __init__(self, artifacts, principal): self.artifacts = artifacts; self.principal = principal
    def validate(self, operation, outcome):
        artifact = self.artifacts.read(str(outcome.result_ref), principal=self.principal)
        accepted = not artifact.expired and artifact.content_sha256 == outcome.result_sha256 and bool(artifact.content)
        return ValidationDecision(accepted, "resume-export-v1", outcome.result_ref, outcome.result_sha256, (str(outcome.result_ref),), False, None if accepted else "export_artifact_invalid")


class _ExportSafetyGate:
    def __init__(self, store, command, principal): self.store=store; self.command=command; self.principal=principal
    def evaluate(self, operation, decision):
        version = self.store.get(ResumeVersion, self.command.resume_version_id, principal=self.principal)
        allowed = version.status == ResumeVersionStatus.CONFIRMED
        return SafetyDecision(allowed, {"confirmed_version": allowed, "restricted_artifact": True}, None if allowed else "export_requires_confirmed_version")


class _ExportCommitter:
    def __init__(self, service, command, settings_hash, principal): self.service=service; self.command=command; self.settings_hash=settings_hash; self.principal=principal
    def commit(self, operation, checkpoint):
        try: record = self.service.store.get(ExportRecord, self.command.export_id, principal=self.principal)
        except ObjectNotFoundError:
            now = self.service.clock()
            record = ExportRecord(export_id=self.command.export_id, resume_version_id=self.command.resume_version_id, format=self.command.format, template_id=self.command.template_id, template_version=self.command.template_version, settings_sha256=self.settings_hash, status=ExportStatus.AVAILABLE, artifact_id=checkpoint.result_ref, content_sha256=checkpoint.result_sha256, revision=1, created_at=now, available_at=now)
            self.service.store.create(record, principal=self.principal, workspace_id=self.command.workspace_id)
        if record.artifact_id != checkpoint.result_ref or record.content_sha256 != checkpoint.result_sha256: raise ExportServiceError("export_id_conflict")
        return CommitReceipt(record.export_id)


class _ReuseValidator:
    def __init__(self, record): self.record=record
    def validate(self, operation, outcome): return ValidationDecision(True, "resume-export-reuse-v1", outcome.result_ref, outcome.result_sha256, (str(outcome.result_ref),))
class _ReuseSafetyGate:
    def evaluate(self, operation, decision): return SafetyDecision(True, {"reused_immutable_export": True})
class _ReuseCommitter:
    def __init__(self, record): self.record=record
    def commit(self, operation, checkpoint): return CommitReceipt(self.record.export_id)
